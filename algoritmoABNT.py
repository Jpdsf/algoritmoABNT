from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL

def estilos(documento):
    estilo1 = documento.styles.add_style('Titulo1', WD_STYLE_TYPE.PARAGRAPH)
    
    estilo1.font.name = 'Times New Roman'
    estilo1.font.size = Pt(12)         
    estilo1.font.bold = True           
    estilo1.paragraph_format.space_after = Pt(8)        
    estilo1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    estilo1.font.color.rgb = RGBColor(0, 0, 0)

    estilo2 = documento.styles.add_style('Titulo2', WD_STYLE_TYPE.PARAGRAPH)
    
    estilo2.font.name = 'Times New Roman'
    estilo2.font.size = Pt(12)          
    estilo2.font.bold = False
    estilo2.paragraph_format.space_after = Pt(8)        
    estilo2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    estilo2.font.color.rgb = RGBColor(0, 0, 0)

    estilo3 = documento.styles.add_style('Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    estilo3.font.name = 'Times New Roman'
    estilo3.font.size = Pt(12)          
    estilo3.paragraph_format.line_spacing = 1.5
    estilo3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    estilo3.font.color.rgb = RGBColor(0, 0, 0)

def formatar_tabela(tabela):
    for linha_idx, linha in enumerate(tabela.rows):
        for celula_idx, celula in enumerate(linha.cells):
            for paragrafo in celula.paragraphs:
                for run in paragrafo.runs:
                    run.font.size = Pt(12)  
                    run.font.bold = False  
                    run.font.line_spacing = 1.5 
                    run.font.color.rgb = RGBColor(0, 0, 0)  
                    run.font.name = 'Times New Roman'
                    run.font.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            definir_bordas(celula, linha_idx, celula_idx, len(tabela.rows), len(linha.cells))

def definir_bordas(celula, linha_idx, celula_idx, total_linhas, total_colunas):
    tc_pr = celula._element.get_or_add_tcPr()

    esquerda = OxmlElement('w:left')
    direita = OxmlElement('w:right')
    top = OxmlElement('w:top')
    buttom = OxmlElement('w:buttom')

    esquerda.set(qn('w:val'), 'nil')
    direita.set(qn('w:val'), 'nil')
    top.set(qn('w:val'), 'nil')
    buttom.set(qn('w:val'), 'nil')
    tc_pr.append(esquerda)
    tc_pr.append(direita)
    tc_pr.append(top)
    tc_pr.append(buttom)

    if linha_idx > 0 and linha_idx < total_linhas - 1:
        superior = OxmlElement('w:top')
        inferior = OxmlElement('w:bottom')
        superior.set(qn('w:val'), 'nil')
        inferior.set(qn('w:val'), 'nil')
        tc_pr.append(superior)
        tc_pr.append(inferior)
    else:
        if linha_idx == 0:  
            borda = OxmlElement('w:top')
            borda.set(qn('w:val'), 'single')
            borda.set(qn('w:sz'), '12') 
            borda.set(qn('w:color'), '000000')
            tc_pr.append(borda)
        elif linha_idx == total_linhas - 1:  
            borda = OxmlElement('w:bottom')
            borda.set(qn('w:val'), 'single')
            borda.set(qn('w:sz'), '12')
            borda.set(qn('w:color'), '000000')
            tc_pr.append(borda)
        


def formatar_documento(caminho_entrada, caminho_saida):
    doc = Document(caminho_entrada)
    
    estilos(doc)
    
    for tabela in doc.tables:
        formatar_tabela(tabela)
    
    for paragrafo in doc.paragraphs:
        texto = paragrafo.text
        
        if '<titulo1>' in texto and '</titulo1>' in texto:
            inicio = texto.find('<titulo1>') + len('<titulo1>')
            fim = texto.find('</titulo1>')
            conteudo = texto[inicio:fim].strip().upper()
            paragrafo.text = conteudo
            paragrafo.style = doc.styles['Titulo1']
            
        elif '<titulo2>' in texto and '</titulo2>' in texto:
            inicio = texto.find('<titulo2>') + len('<titulo2>')
            fim = texto.find('</titulo2>')
            conteudo = texto[inicio:fim].strip().upper()
            paragrafo.text = conteudo
            paragrafo.style = doc.styles['Titulo2']

        elif '<p>' in texto and '</p>' in texto:
            inicio = texto.find('<p>') + len('<p>')
            fim = texto.find('</p>')
            conteudo = texto[inicio:fim].strip()
            paragrafo.text = conteudo
            paragrafo.style = doc.styles['Paragraph']

    doc.save(caminho_saida)

caminho_entrada = "D:\\Projetos\\algoritmoABNT\\arquivo_entrada.docx"
caminho_saida = "arquivo_saida.docx"

formatar_documento(caminho_entrada, caminho_saida)
